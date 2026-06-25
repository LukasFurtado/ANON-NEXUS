from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".csv", ".xlsx", ".xlsm"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Formato ainda nao suportado: {suffix}")

    if suffix in {".txt", ".csv"}:
        return _read_text_file(path)

    if suffix == ".rtf":
        return _extract_rtf(path)

    if suffix == ".docx":
        return _extract_docx(path)

    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)

    if suffix == ".pdf":
        return _extract_pdf(path)

    if suffix == ".doc":
        raise ValueError("DOC legado requer LibreOffice ou conversor local configurado.")

    return ""


def inspect_source_document(path: Path, extracted_text: str) -> dict[str, object]:
    suffix = path.suffix.lower()
    return {
        "extension": suffix,
        "size_bytes": path.stat().st_size if path.exists() else 0,
        "text_characters": len(extracted_text),
        "text_lines": len(extracted_text.splitlines()),
        "likely_delimiter": _likely_delimiter(extracted_text) if suffix in {".csv", ".txt"} else "",
        "has_tabular_text": _has_tabular_text(extracted_text),
        "ocr_recommended": len(_alnum_text(extracted_text)) < 80,
    }


def _extract_pdf(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("PyMuPDF nao esta instalado no ambiente local.") from exc

    pages: list[str] = []
    with fitz.open(path) as document:
        for page in document:
            page_text = _extract_pdf_page_structured(page)
            pages.append(page_text)
    return "\n".join(pages)


def _extract_pdf_page_structured(page) -> str:
    parts: list[str] = []
    try:
        tables = page.find_tables()
        for table in tables:
            extracted = table.extract()
            if extracted:
                parts.append("\n".join("\t".join(str(cell or "") for cell in row) for row in extracted))
    except Exception:
        pass

    try:
        blocks = page.get_text("blocks", sort=True)
        lines = [
            str(block[4]).strip()
            for block in blocks
            if len(block) >= 5 and str(block[4]).strip()
        ]
        if lines:
            parts.append("\n".join(lines))
    except Exception:
        text = page.get_text("text", sort=True)
        if text.strip():
            parts.append(text.strip())

    if not parts:
        parts.append(page.get_text("text"))
    return "\n".join(parts)


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _likely_delimiter(text: str) -> str:
    sample = [line for line in text.splitlines()[:10] if line.strip()]
    if not sample:
        return ""
    counts = {";": 0, ",": 0, "\t": 0, "|": 0}
    for line in sample:
        for delimiter in counts:
            counts[delimiter] += line.count(delimiter)
    delimiter, count = max(counts.items(), key=lambda item: item[1])
    return delimiter if count else ""


def _has_tabular_text(text: str) -> bool:
    lines = [line for line in text.splitlines()[:20] if line.strip()]
    if not lines:
        return False
    return any(line.count(";") >= 2 or line.count("\t") >= 2 for line in lines)


def _alnum_text(text: str) -> str:
    return "".join(char for char in text if char.isalnum())


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx nao esta instalado no ambiente local.") from exc

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("openpyxl nao esta instalado no ambiente local para leitura de Excel.") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for worksheet in workbook.worksheets:
        parts.append(f"# ABA: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell) for cell in row]
            if any(value.strip() for value in values):
                parts.append("\t".join(values).rstrip())
    workbook.close()
    return "\n".join(parts)


def _extract_rtf(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    text = raw.replace("\\par", "\n")
    cleaned: list[str] = []
    skip = False
    for char in text:
        if char == "{":
            skip = True
            continue
        if char == "}":
            skip = False
            continue
        if not skip and char != "\\":
            cleaned.append(char)
    return "".join(cleaned)
