from datetime import datetime
import json
from pathlib import Path
import re
from textwrap import wrap
from typing import Any

from app.models.schemas import DocumentKind


PROJECT_ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_DIR = PROJECT_ROOT / "resources" / "templates"
DOCX_TEMPLATE = TEMPLATE_DIR / "modelo_raf.docx"
LOGO_PATH = TEMPLATE_DIR / "logo_pcpe_header.png"


def export_text(
    job_id: str,
    anonymized_text: str,
    original_filename: str = "documento",
    metadata: dict[str, Any] | None = None,
) -> dict[str, str]:
    export_dir = Path("data") / "exports" / job_id
    export_dir.mkdir(parents=True, exist_ok=True)

    txt_path = export_dir / "anonimizado.txt"
    docx_path = export_dir / "anonimizado.docx"
    pdf_path = export_dir / "anonimizado.pdf"
    warnings_path = export_dir / "avisos.pdf"

    metadata = metadata or {}
    is_pdf_bank_statement = (
        metadata.get("document_kind") == DocumentKind.extrato_bancario.value
        and Path(original_filename).suffix.lower() == ".pdf"
    )
    if is_pdf_bank_statement:
        _export_pdf(pdf_path, anonymized_text, original_filename, metadata)
        exports = {"pdf": str(pdf_path)}
        if _validation_warnings(metadata):
            _export_warnings_pdf(warnings_path, original_filename, metadata)
            exports["avisos"] = str(warnings_path)
        return exports

    _export_txt(txt_path, anonymized_text, original_filename, metadata)
    _export_docx(docx_path, anonymized_text, original_filename, metadata)
    _export_pdf(pdf_path, anonymized_text, original_filename, metadata)

    exports = {"txt": str(txt_path), "docx": str(docx_path), "pdf": str(pdf_path)}
    if Path(original_filename).suffix.lower() == ".csv":
        csv_path = export_dir / "anonimizado.csv"
        _export_csv(csv_path, anonymized_text)
        exports["csv"] = str(csv_path)
    if _validation_warnings(metadata):
        _export_warnings_pdf(warnings_path, original_filename, metadata)
        exports["avisos"] = str(warnings_path)
    control_path = export_dir / "controle_interno.pdf"
    if _control_rows(metadata):
        _export_control_pdf(control_path, original_filename, metadata)
        exports["controle"] = str(control_path)

    return exports


def export_audit_manifest(
    job_id: str,
    original_filename: str,
    metadata: dict[str, Any],
    export_hashes: dict[str, str],
) -> str:
    export_dir = Path("data") / "exports" / job_id
    export_dir.mkdir(parents=True, exist_ok=True)
    audit_path = export_dir / "auditoria_interna.json"
    products = []
    for format_name, sha256 in sorted(export_hashes.items()):
        products.append(
            {
                "format": format_name,
                "sha256": sha256,
                "usage_classification": _product_usage_classification(format_name),
                "contains_original_values": format_name == "controle",
                "external_sharing_allowed": format_name in _external_product_formats(),
            }
        )
    payload = {
        "schema": "ANON-AUDITORIA-INTERNA-v1",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "request_title": metadata.get("request_title") or "Nao informada",
        "original_filename": original_filename,
        "document_kind": metadata.get("document_kind") or "Nao informado",
        "model": metadata.get("model") or "Nao informado",
        "anon_version": metadata.get("anon_version") or "Nao informada",
        "source_sha256": metadata.get("source_sha256") or "Nao calculado",
        "processing": {
            "time_seconds": metadata.get("processing_time_seconds"),
            "ocr_used": bool(metadata.get("ocr_used")),
            "structure_preserved": bool(metadata.get("structure_preserved")),
            "validation_status": metadata.get("validation_status") or "Nao informada",
        },
        "counts": {
            "entities_found": int(metadata.get("entities_found") or 0),
            "replacements_applied": int(metadata.get("replacements_applied") or 0),
            "validation_warnings": len(_validation_warnings(metadata)),
            "control_rows": len(_control_rows(metadata)),
        },
        "ollama_json": metadata.get("ollama_metrics") or {},
        "post_validation": metadata.get("post_validation") or {},
        "confidence": metadata.get("confidence") or {},
        "human_review": {
            "mode": "enabled",
            "status": "pending_operator_review" if metadata.get("review_items") else "no_items",
            "items": metadata.get("review_items") or [],
        },
        "communication_summary": metadata.get("communication_summary") or {},
        "nce_context": metadata.get("nce_context") or {},
        "nce_file_context": metadata.get("nce_file_context") or {},
        "product_policy": {
            "external_products": sorted(_external_product_formats()),
            "internal_restricted_products": sorted(_internal_product_formats()),
            "note": (
                "Produtos externos sao destinados a uso apos revisao humana. Arquivos internos restritos "
                "servem para auditoria, conferencia, avisos e cadeia de custodia."
            ),
        },
        "products": products,
    }
    audit_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(audit_path)


def _external_product_formats() -> set[str]:
    return {"txt", "docx", "pdf", "csv"}


def _internal_product_formats() -> set[str]:
    return {"avisos", "controle", "auditoria"}


def _product_usage_classification(format_name: str) -> str:
    if format_name in _external_product_formats():
        return "produto_externo_revisao_obrigatoria"
    if format_name == "controle":
        return "interno_restrito_tabela_de_correspondencia"
    if format_name == "avisos":
        return "interno_restrito_avisos_de_validacao"
    if format_name == "auditoria":
        return "interno_restrito_manifesto_de_auditoria"
    return "interno_restrito"


def _summary_lines(original_filename: str, metadata: dict[str, Any]) -> list[str]:
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    communication_summary = metadata.get("communication_summary") if isinstance(metadata.get("communication_summary"), dict) else {}
    lines = [
        "RESUMO OPERACIONAL",
        f"Solicitacao: {metadata.get('request_title') or 'Nao informada'}",
        f"Arquivo de origem: {original_filename}",
        f"Perfil documental: {metadata.get('document_kind') or 'Nao informado'}",
        f"Modelo: {metadata.get('model') or 'Nao informado'}",
        f"Versao do ANON: {metadata.get('anon_version') or 'Nao informada'}",
        f"Tempo de processamento: {metadata.get('processing_time_seconds', 'Nao informado')} s",
        f"OCR: {'Utilizado' if metadata.get('ocr_used') else 'Nao utilizado'}",
        f"Estrutura: {'Preservada' if metadata.get('structure_preserved') else 'Nao preservada'}",
        f"Validacao: {metadata.get('validation_status') or 'Nao informada'}",
        f"Hash SHA-256 original: {metadata.get('source_sha256') or 'Nao calculado'}",
        f"Entidades identificadas: {metadata.get('entities_found', 0)}",
        f"Substituicoes aplicadas: {metadata.get('replacements_applied', 0)}",
        f"Eventos internos de comunicacao: {communication_summary.get('events', 0)}",
        f"Ultimo estagio interno: {communication_summary.get('last_stage') or 'Nao informado'}",
        f"Gerado em: {generated_at}",
    ]
    protection = _data_protection(metadata)
    if protection.get("notice"):
        lines.append(str(protection["notice"]))
    return lines


def _data_protection(metadata: dict[str, Any]) -> dict[str, str]:
    value = metadata.get("data_protection")
    return value if isinstance(value, dict) else {}


def _validation_warnings(metadata: dict[str, Any]) -> list[str]:
    warnings = metadata.get("validation_warnings") or []
    return [str(warning) for warning in warnings if str(warning).strip()] if isinstance(warnings, list) else []


def _warning_explanation(warning: str) -> str:
    normalized = warning.lower()
    if "ia local foi acionada" in normalized and "nao retornou resposta aproveitavel" in normalized:
        return (
            "Justificativa: a IA local foi solicitada, mas nao devolveu uma resposta utilizavel para auditoria automatica "
            "nesta execucao. O ANON manteve a entrega com regras locais de apoio e sinalizou a necessidade de revisao humana."
        )
    if "nao retornou entidades aproveitaveis em json" in normalized:
        return (
            "Justificativa: o modelo local foi acionado e respondeu, mas a resposta nao apresentou lista JSON de "
            "entidades suficientemente estruturada para auditoria automatica. O ANON preservou o fluxo com regras "
            "locais de apoio e registrou este aviso para revisao humana."
        )
    if "termos protegidos do perfil" in normalized:
        return "Justificativa: possivel alteracao em termo protegido do perfil documental; revisar estrutura, colunas, titulos e expressoes tecnicas."
    if "valores possivelmente alterados" in normalized:
        return "Justificativa: possivel divergencia em valores monetarios; conferir quantias e formatacao financeira."
    if "datas possivelmente alteradas" in normalized:
        return "Justificativa: possivel divergencia em datas; conferir datas de registro, emissao, movimentacao ou periodo."
    return "Justificativa: aviso automatico de validacao; revisar o ponto indicado antes de uso externo ou oficial."


def _control_rows(metadata: dict[str, Any]) -> list[Any]:
    rows = metadata.get("control_table") or []
    return rows if isinstance(rows, list) else []


def _row_value(row: Any, key: str) -> Any:
    if isinstance(row, dict):
        return row.get(key, "")
    return getattr(row, key, "")


def _compact_cell(value: Any, limit: int = 220) -> str:
    text = " ".join(str(value).split())
    if len(text) <= limit:
        return text
    return f"{text[: limit - 15].rstrip()}... [truncado]"


def _export_txt(path: Path, text: str, original_filename: str, metadata: dict[str, Any]) -> None:
    lines = [
        "NEXUS ANON - DOCUMENTO ANONIMIZADO",
        "POLICIA CIVIL DO ESTADO DE PERNAMBUCO",
        "-" * 78,
        "",
        *_summary_lines(original_filename, metadata),
        "",
        "-" * 78,
        "PRODUTO FINAL DA ANONIMIZACAO DE DADOS",
        "O conteudo abaixo corresponde ao resultado desidentificado gerado pelo NEXUS ANON.",
        "-" * 78,
        "",
    ]
    for paragraph in text.splitlines():
        if not paragraph.strip():
            lines.append("")
            continue
        if _looks_like_delimited_row(paragraph):
            lines.append(paragraph)
            continue
        lines.extend(wrap(paragraph, width=110, replace_whitespace=False, drop_whitespace=False) or [""])
    lines.extend(["", "-" * 78, "FIM DO PRODUTO FINAL DA ANONIMIZACAO DE DADOS"])
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _looks_like_delimited_row(value: str) -> bool:
    return value.count(";") >= 2 or value.count(",") >= 5 or "\t" in value


def _export_csv(path: Path, text: str) -> None:
    path.write_text(text.rstrip("\r\n") + "\n", encoding="utf-8-sig")


def _export_docx(path: Path, text: str, original_filename: str, metadata: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        path.write_text(text, encoding="utf-8")
        return

    document = Document(DOCX_TEMPLATE) if DOCX_TEMPLATE.exists() else Document()
    _clear_docx_body_preserving_section(document)
    _apply_docx_source_layout(document, metadata, WD_ORIENT)
    _apply_docx_protection_metadata(document, metadata)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENTO ANONIMIZADO")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("NEXUS ANON - POLICIA CIVIL DO ESTADO DE PERNAMBUCO")
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run("\n".join(_summary_lines(original_filename, metadata)))

    document.add_paragraph("")

    marker = document.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    marker_run = marker.add_run("PRODUTO FINAL DA ANONIMIZACAO DE DADOS")
    marker_run.bold = True
    marker_run.font.name = "Arial"
    marker_run.font.size = Pt(11)

    for line in text.splitlines():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.add_run("FIM DO PRODUTO FINAL DA ANONIMIZACAO DE DADOS").bold = True

    document.save(path)


def _apply_docx_protection_metadata(document, metadata: dict[str, Any]) -> None:
    protection = _data_protection(metadata)
    marker = protection.get("marker")
    if not marker:
        return
    properties = document.core_properties
    properties.category = "Protecao de dados"
    properties.keywords = "ANON; Protecao de dados; Integridade documental"
    properties.comments = f"Protecao de dados institucional ativa: {marker}"


def _clear_docx_body_preserving_section(document) -> None:
    body = document._body._element
    section_properties = None
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            section_properties = child
        body.remove(child)
    if section_properties is not None:
        body.append(section_properties)


def _apply_docx_source_layout(document, metadata: dict[str, Any], wd_orient) -> None:
    if metadata.get("source_pdf_orientation") != "landscape":
        return
    for section in document.sections:
        section.orientation = wd_orient.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def _export_pdf(path: Path, text: str, original_filename: str, metadata: dict[str, Any]) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        path.write_text(text, encoding="utf-8")
        return

    if _try_export_pdf_facsimile(path, original_filename, metadata):
        return

    page_size = landscape(A4) if metadata.get("source_pdf_orientation") == "landscape" else A4

    doc = SimpleDocTemplate(
        str(path),
        pagesize=page_size,
        rightMargin=1.6 * cm if metadata.get("source_pdf_orientation") == "landscape" else 2 * cm,
        leftMargin=1.6 * cm if metadata.get("source_pdf_orientation") == "landscape" else 2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Documento Anonimizado",
        author="NEXUS ANON",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "InstitutionalTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        alignment=1,
        textColor=colors.HexColor("#003B73"),
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "InstitutionalMeta",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        alignment=0,
        textColor=colors.HexColor("#172026"),
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "InstitutionalBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=0,
        firstLineIndent=0,
        spaceAfter=6,
    )

    story = []
    if LOGO_PATH.exists():
        logo = Image(str(LOGO_PATH), width=8.2 * cm, height=2.2 * cm, kind="proportional")
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("DOCUMENTO ANONIMIZADO", title_style))
    story.append(Paragraph("NEXUS ANON - POLICIA CIVIL DO ESTADO DE PERNAMBUCO", title_style))
    story.append(
        Paragraph(
            "<br/>".join(_escape_pdf_text(line) for line in _summary_lines(original_filename, metadata)),
            meta_style,
        )
    )
    story.append(Paragraph("PRODUTO FINAL DA ANONIMIZACAO DE DADOS", title_style))

    for line in text.splitlines():
        story.append(Paragraph(_escape_pdf_text(line) if line.strip() else "&nbsp;", body_style))

    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("FIM DO PRODUTO FINAL DA ANONIMIZACAO DE DADOS", meta_style))

    doc.build(story)
    _apply_pdf_protection_metadata(path, metadata)


def _try_export_pdf_facsimile(path: Path, original_filename: str, metadata: dict[str, Any]) -> bool:
    source_path = Path(str(metadata.get("source_path") or ""))
    rows = _control_rows(metadata)
    if Path(original_filename).suffix.lower() != ".pdf" or not source_path.exists() or not rows:
        return False

    try:
        import fitz
    except ImportError:
        return False

    try:
        document = fitz.open(source_path)
        total_hits = 0
        replacements = sorted(
            (
                (str(_row_value(row, "original_value")), str(_row_value(row, "anonymous_id")))
                for row in rows
                if str(_row_value(row, "original_value")).strip() and str(_row_value(row, "anonymous_id")).strip()
            ),
            key=lambda item: len(item[0]),
            reverse=True,
        )

        for page in document:
            page_hits = 0
            overlays: list[tuple[Any, str]] = []
            for original, anonymous_id in replacements:
                if _should_skip_pdf_facsimile_replacement(original, anonymous_id, metadata):
                    continue
                rects = _search_pdf_replacement_rects(page, original)
                if not rects:
                    continue
                for rect in rects:
                    expanded = _expanded_pdf_rect(page, rect, anonymous_id)
                    page.add_redact_annot(expanded, fill=(1, 1, 1))
                    overlays.append((expanded, anonymous_id))
                    page_hits += 1
            if page_hits:
                page.apply_redactions()
                for expanded, anonymous_id in overlays:
                    _insert_pdf_marker(page, expanded, anonymous_id)
                total_hits += page_hits

        if total_hits == 0:
            document.close()
            return False

        document.save(path, garbage=4, deflate=True)
        document.close()
        _apply_pdf_protection_metadata(path, metadata)
        return True
    except Exception:
        try:
            if "document" in locals():
                document.close()
        except Exception:
            pass
        return False


def _search_pdf_replacement_rects(page: Any, original: str) -> list[Any]:
    seen: set[tuple[int, int, int, int]] = set()
    rects: list[Any] = []
    for candidate in _pdf_search_candidates(original):
        try:
            found = page.search_for(candidate)
        except Exception:
            found = []
        for rect in found:
            key = tuple(round(float(value) * 10) for value in (rect.x0, rect.y0, rect.x1, rect.y1))
            if key in seen:
                continue
            seen.add(key)
            rects.append(rect)
    return rects


def _pdf_search_candidates(original: str) -> list[str]:
    value = " ".join(str(original or "").split())
    if not value:
        return []
    candidates = [value]
    digits = re.sub(r"\D", "", value)
    if digits:
        candidates.append(digits)
        if len(digits) == 11:
            candidates.extend(
                [
                    f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}",
                    f"{digits[:3]} {digits[3:6]} {digits[6:9]} {digits[9:]}",
                ]
            )
        elif len(digits) == 14:
            candidates.extend(
                [
                    f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}",
                    f"{digits[:2]} {digits[2:5]} {digits[5:8]} {digits[8:12]} {digits[12:]}",
                ]
            )
    candidates.extend([value.upper(), value.title()])
    compacted = re.sub(r"\s+", " ", value.replace("\t", " ")).strip()
    candidates.append(compacted)
    output: list[str] = []
    for candidate in candidates:
        if candidate and candidate not in output:
            output.append(candidate)
    return output


def _should_skip_pdf_facsimile_replacement(original: str, anonymous_id: str, metadata: dict[str, Any]) -> bool:
    if metadata.get("document_kind") != DocumentKind.extrato_bancario.value:
        return False
    value = str(original or "").strip()
    marker = str(anonymous_id or "").strip()
    digits = re.sub(r"\D", "", value)
    if not value or not marker:
        return True
    if re.fullmatch(r"[\d\s./-]+", value) and 1 <= len(digits) <= 10:
        return True
    if marker.startswith("[IDENTIFICADOR_") and digits:
        return True
    return False


def _expanded_pdf_rect(page: Any, rect: Any, anonymous_id: str) -> Any:
    try:
        import fitz
    except ImportError:
        return rect
    page_rect = page.rect
    width_needed = max(float(rect.width), len(anonymous_id) * max(float(rect.height), 6) * 0.42)
    extra = max(0, width_needed - float(rect.width))
    return fitz.Rect(
        max(float(page_rect.x0), float(rect.x0) - min(extra * 0.2, 8)),
        max(float(page_rect.y0), float(rect.y0) - 0.8),
        min(float(page_rect.x1), float(rect.x1) + extra + 3),
        min(float(page_rect.y1), float(rect.y1) + 1.2),
    )


def _insert_pdf_marker(page: Any, rect: Any, anonymous_id: str) -> None:
    fontsize = _pdf_marker_font_size(rect, anonymous_id)
    rc = page.insert_textbox(
        rect,
        anonymous_id,
        fontsize=fontsize,
        fontname="helv",
        color=(0, 0, 0),
        align=0,
    )
    if rc >= 0:
        return
    try:
        import fitz
    except ImportError:
        return
    expanded = fitz.Rect(rect.x0, rect.y0, min(page.rect.x1, rect.x1 + 90), min(page.rect.y1, rect.y1 + 4))
    for fallback_size in (max(4.0, fontsize - 0.8), 3.8, 3.4):
        rc = page.insert_textbox(
            expanded,
            anonymous_id,
            fontsize=fallback_size,
            fontname="helv",
            color=(0, 0, 0),
            align=0,
        )
        if rc >= 0:
            return
    page.insert_text((expanded.x0, min(expanded.y1 - 1, expanded.y0 + 6)), anonymous_id, fontsize=3.4, fontname="helv", color=(0, 0, 0))


def _pdf_marker_font_size(rect: Any, anonymous_id: str) -> float:
    if not anonymous_id:
        return 6
    by_height = max(4.2, min(8.0, float(rect.height) * 0.64))
    by_width = max(4.2, min(8.0, float(rect.width) / max(len(anonymous_id) * 0.52, 1)))
    return round(min(by_height, by_width), 2)


def _export_control_pdf(path: Path, original_filename: str, metadata: dict[str, Any]) -> None:
    rows = _control_rows(metadata)
    if not rows:
        return
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    page_size = landscape(A4)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=page_size,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="Controle Interno de Anonimizacao - ANON",
        author="ANON",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ControlTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=13,
        alignment=1,
        textColor=colors.HexColor("#003B73"),
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "ControlBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        alignment=0,
        textColor=colors.HexColor("#172026"),
        spaceAfter=6,
    )

    story: list[Any] = [
        Paragraph("TABELA DE CONTROLE DE ANONIMIZACAO - USO INTERNO", title_style),
        Paragraph(
            "Arquivo separado do produto final. Contem valores originais e deve permanecer restrito a auditoria, conferencia, revisao interna e cadeia de custodia.",
            body_style,
        ),
        Paragraph("<br/>".join(_escape_pdf_text(line) for line in _summary_lines(original_filename, metadata)), body_style),
        Spacer(1, 0.25 * cm),
    ]
    table_data = [["Valor Original", "Tipo", "Identificador", "Ocorrencias"]]
    for row in rows:
        table_data.append(
            [
                Paragraph(_escape_pdf_text(_compact_cell(_row_value(row, "original_value"), 180)), body_style),
                Paragraph(_escape_pdf_text(_compact_cell(_row_value(row, "entity_type"), 50)), body_style),
                Paragraph(_escape_pdf_text(_compact_cell(_row_value(row, "anonymous_id"), 50)), body_style),
                str(_row_value(row, "occurrences")),
            ]
        )
    control_table = Table(table_data, colWidths=[10.5 * cm, 3.3 * cm, 4.0 * cm, 2.2 * cm], repeatRows=1)
    control_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003B73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C3CF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(control_table)
    doc.build(story)
    _apply_pdf_protection_metadata(path, metadata)


def _export_warnings_pdf(path: Path, original_filename: str, metadata: dict[str, Any]) -> None:
    warnings = _validation_warnings(metadata)
    if not warnings:
        return
    from app.services.knowledge_base import json_contract_prompt
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title="Avisos de Validacao - ANON",
        author="ANON",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "WarningsTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        alignment=1,
        textColor=colors.HexColor("#003B73"),
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "WarningsBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        alignment=0,
        spaceAfter=7,
    )

    metrics = metadata.get("ollama_metrics") or {}
    story: list[Any] = [
        Paragraph("AVISOS DE VALIDACAO - REVISAO OBRIGATORIA", title_style),
        Paragraph("Documento complementar de auditoria operacional do ANON.", body_style),
        Spacer(1, 0.2 * cm),
        Paragraph("<br/>".join(_escape_pdf_text(line) for line in _summary_lines(original_filename, metadata)), body_style),
        Spacer(1, 0.25 * cm),
        Paragraph("Avaliacao da resposta JSON da IA local", title_style),
        Paragraph(
            _escape_pdf_text(
                "Blocos enviados a IA: "
                f"{metrics.get('chunks_processed', 0)} | "
                "Blocos recusados por JSON nao aproveitavel: "
                f"{metrics.get('json_rejected_chunks', 0)} | "
                "Tentativas de correcao: "
                f"{metrics.get('correction_attempts', 0)} | "
                "Correcoes aproveitadas: "
                f"{metrics.get('correction_successes', 0)}"
            ),
            body_style,
        ),
        Paragraph(
            _escape_pdf_text(_format_json_rejection_reasons(metrics)).replace("\n", "<br/>"),
            body_style,
        ),
        Paragraph(_escape_pdf_text(json_contract_prompt()).replace("\n", "<br/>"), body_style),
        Spacer(1, 0.25 * cm),
        Paragraph("Avisos registrados", title_style),
    ]

    for index, warning in enumerate(warnings, start=1):
        text = f"{index}. {warning}<br/>{_warning_explanation(warning)}"
        story.append(Paragraph(_escape_pdf_text(text).replace("&lt;br/&gt;", "<br/>"), body_style))
    story.extend(
        [
            Spacer(1, 0.25 * cm),
            Paragraph("Consideracoes finais", title_style),
            Paragraph(
                "Os avisos nao invalidam automaticamente o produto anonimizado, mas indicam pontos que exigem revisao humana antes de uso externo, juntada formal ou compartilhamento institucional. Quando houver JSON nao aproveitavel da IA local, o ANON registra a ocorrencia, preserva o processamento com regras locais de apoio e destaca a necessidade de conferencia qualificada.",
                body_style,
            ),
        ]
    )
    doc.build(story)
    _apply_pdf_protection_metadata(path, metadata)


def _format_json_rejection_reasons(metrics: dict) -> str:
    reasons = metrics.get("json_rejection_reasons") or []
    if not reasons:
        return "Motivos de recusa JSON: nenhum bloco foi recusado ou todos os blocos ficaram dentro do contrato esperado."
    unique_reasons: list[str] = []
    for reason in reasons:
        text = str(reason)
        if text not in unique_reasons:
            unique_reasons.append(text)
    return "Motivos de recusa JSON:\n" + "\n".join(f"- {reason}" for reason in unique_reasons[:6])


def _apply_pdf_protection_metadata(path: Path, metadata: dict[str, Any]) -> None:
    protection = _data_protection(metadata)
    marker = protection.get("marker")
    if not marker or not path.exists():
        return
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return

    temp_path = path.with_suffix(".metadata.tmp.pdf")
    try:
        reader = PdfReader(str(path))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        existing = reader.metadata or {}
        writer.add_metadata(
            {
                **{str(key): str(value) for key, value in existing.items() if value is not None},
                "/Subject": "Protecao de dados institucional",
                "/Keywords": f"ANON; Protecao de dados; {marker}",
                "/Creator": "ANON - Protecao de dados",
                "/Producer": "ANON",
            }
        )
        with temp_path.open("wb") as file:
            writer.write(file)
        temp_path.replace(path)
    except Exception:
        temp_path.unlink(missing_ok=True)


def _escape_pdf_text(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_processing_log(group_id: str, metadata: dict[str, Any]) -> str:
    export_dir = Path("data") / "exports" / group_id
    export_dir.mkdir(parents=True, exist_ok=True)
    log_path = export_dir / "log_processamento.pdf"
    _export_log_pdf(log_path, metadata)
    return str(log_path)


def _export_log_pdf(path: Path, metadata: dict[str, Any]) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        lines = ["NEXUS ANON - LOG DE PROCESSAMENTO", *metadata.get("summary_lines", [])]
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
        title="Log de Processamento - NEXUS ANON",
        author="NEXUS ANON",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "LogTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        alignment=1,
        textColor=colors.HexColor("#003B73"),
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "LogBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        alignment=0,
        spaceAfter=4,
    )

    story = []
    if LOGO_PATH.exists():
        logo = Image(str(LOGO_PATH), width=8.2 * cm, height=2.2 * cm, kind="proportional")
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("LOG DE PROCESSAMENTO E CADEIA DE CUSTODIA", title_style))
    for line in metadata.get("summary_lines", []):
        story.append(Paragraph(_escape_pdf_text(str(line)), body_style))

    file_rows = metadata.get("files", [])
    if file_rows:
        story.append(Spacer(1, 0.25 * cm))
        story.append(Paragraph("Arquivos e hashes registrados", title_style))
        table_data = [["Arquivo", "Versao", "Hash original", "TXT", "DOCX", "PDF", "CSV/Avisos/Controle"]]
        for item in file_rows:
            table_data.append(
                [
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("filename", ""), 80)), body_style),
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("anon_version", ""), 24)), body_style),
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("source_sha256", ""), 72)), body_style),
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("txt_sha256", ""), 72)), body_style),
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("docx_sha256", ""), 72)), body_style),
                    Paragraph(_escape_pdf_text(_compact_cell(item.get("pdf_sha256", ""), 72)), body_style),
                    Paragraph(
                        _escape_pdf_text(
                            _compact_cell(
                                item.get("csv_sha256") or item.get("avisos_sha256") or item.get("controle_sha256", ""),
                                72,
                            )
                        ),
                        body_style,
                    ),
                ]
            )
        table = Table(table_data, colWidths=[3.0 * cm, 1.4 * cm, 3.2 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003B73")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C3CF")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)

    doc.build(story)
